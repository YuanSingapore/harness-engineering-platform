import os
import datetime
import json
from typing import List
import docx
import anthropic
from models import MCQQuestion, WordOut
from dotenv import load_dotenv

load_dotenv()

APP_START_DATE = "2026-08-02"
NUM_CHAPTERS = 5

anthropic_client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))


def get_word_bank_dir() -> str:
    return os.path.join(os.path.dirname(__file__), "../../../word-bank")


def get_today_chapter_number() -> int:
    today = datetime.date.today()
    start = datetime.date.fromisoformat(APP_START_DATE)
    day_index = (today - start).days
    return (day_index % NUM_CHAPTERS) + 1


def parse_chapter_docx(chapter_num: int) -> List[MCQQuestion]:
    path = os.path.join(
        get_word_bank_dir(),
        f"P4_Vocabulary_MCQ_Workbook_Chapter_{chapter_num:02d}.docx"
    )
    doc = docx.Document(path)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    questions = []
    i = 0
    while i < len(texts):
        # Detect start of a question block: "N. Word"
        parts = texts[i].split(". ", 1)
        if len(parts) == 2 and parts[0].isdigit():
            word = parts[1].strip()
            # Skip: Part of Speech, Category, Meaning, Synonym, "Choose the most..."
            i += 6
            if i + 6 >= len(texts):
                break
            question_text = texts[i]; i += 1
            choice_a = texts[i].removeprefix("A. "); i += 1
            choice_b = texts[i].removeprefix("B. "); i += 1
            choice_c = texts[i].removeprefix("C. "); i += 1
            choice_d = texts[i].removeprefix("D. "); i += 1
            answer_line = texts[i]; i += 1  # "Answer: B"
            explanation_line = texts[i]; i += 1  # "Explanation: ..."
            # Skip "Composition Example: ..."
            if i < len(texts) and texts[i].startswith("Composition Example"):
                i += 1

            letter = answer_line.replace("Answer:", "").strip()
            choice_map = {"A": choice_a, "B": choice_b, "C": choice_c, "D": choice_d}
            correct = choice_map.get(letter, choice_a)
            explanation = explanation_line.replace("Explanation:", "").strip()

            questions.append(MCQQuestion(
                word=word,
                question=question_text,
                choices=[choice_a, choice_b, choice_c, choice_d],
                correct_answer=correct,
                explanation=explanation,
            ))
        else:
            i += 1

    return questions


def generate_mcq_round(words: List[WordOut], previous_questions: List[str]) -> List[MCQQuestion]:
    word_list = "\n".join([
        f"- {w.word} ({w.part_of_speech}): {w.meaning}"
        for w in words
    ])
    prev = "\n".join(previous_questions) if previous_questions else "None"

    prompt = f"""You are a friendly English vocabulary tutor for a Singapore Primary 4 student.

For each word below, generate a multiple-choice question where:
- The question gives the MEANING of the word (do not use the word itself in the question)
- The student must choose the correct WORD from 4 options
- The 3 wrong options are plausible words from the same category as the correct word
- Use encouraging, positive language

Words to quiz:
{word_list}

Previously used questions (DO NOT repeat these):
{prev}

Return ONLY a JSON array. Each object must have exactly these fields:
{{
    "word": "the correct word",
    "question": "the question text showing the meaning",
    "choices": ["word1", "word2", "word3", "word4"],
    "correct_answer": "the correct word verbatim from choices",
    "explanation": "brief encouraging explanation"
}}"""

    response = anthropic_client.messages.create(
        model="rsn.claude-sonnet-4-6",
        max_tokens=2048,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.content[0].text.strip()
    data = json.loads(raw)
    return [MCQQuestion(**q) for q in data]


def write_topup_docx(wrong_words: List[str], date: str):
    path = os.path.join(get_word_bank_dir(), "P4_Vocabulary_MCQ_Workbook_topup_words.docx")
    if os.path.exists(path):
        doc = docx.Document(path)
    else:
        doc = docx.Document()
        doc.add_heading("MCQ Workbook Wrong Words Log", 0)

    doc.add_paragraph(f"--- {date} ---")
    for word in wrong_words:
        doc.add_paragraph(f"• {word}")
    doc.save(path)
